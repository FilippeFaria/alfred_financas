from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "documentos"
OUTPUT_FILE = OUTPUT_DIR / "Alfred_Mapa_de_Documentacao_do_Produto.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for m, val in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, font_name="Aptos", size=11, bold=False, color=None, italic=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", *, style=None, align=None, spacing_after=0, spacing_before=0):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        format_run(run)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(spacing_after)
    pf.space_before = Pt(spacing_before)
    pf.line_spacing = 1.12
    return p


def add_shaded_callout(doc, title, body_lines, fill="EAF3F8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.2)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(title)
    format_run(r, size=11.5, bold=True, color="123E4A")

    for line in body_lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(line)
        format_run(r, size=10.2, color="203040")

    return table


def build_document():
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(2.1)

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(11)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ALFRED FINANÇAS")
    format_run(r, size=13, bold=True, color="5E7B8A")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Mapa de Documentação do Produto")
    format_run(r, size=25, bold=True, color="123E4A")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(
        "Base inicial para organizar documentos do Alfred por tema, manter decisões registradas "
        "e facilitar a evolução do produto com clareza e rastreabilidade."
    )
    format_run(r, size=11.5, color="3C4D56")
    p.paragraph_format.line_spacing = 1.25

    add_shaded_callout(
        doc,
        "Objetivo desta biblioteca",
        [
            "Criar um conjunto de documentos vivos para diferentes disciplinas do produto.",
            "Cada arquivo deve responder uma pergunta prática: o que foi decidido, por quê e qual impacto isso gera.",
            "A proposta aqui já separa os cinco eixos pedidos: Design, LGPD/Regulamentação, Engenharia de Dados, Ciência de Dados e Contabilidade.",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Versão 1.0  |  {date.today():%d/%m/%Y}")
    format_run(r, size=10.5, color="6A7A85")

    doc.add_page_break()

    # Section 1
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run("1. Como usar esta documentação")
    format_run(r, size=17, bold=True, color="123E4A")

    for text in [
        "Este documento funciona como um mapa de referência. A ideia não é resolver tudo de uma vez, mas organizar a produção dos próximos arquivos do produto.",
        "Cada área deve ter um documento próprio, com escopo claro e uma linguagem simples o bastante para apoiar decisões reais de trabalho.",
        "Sempre que uma decisão afetar produto, dados, compliance ou contabilidade, vale registrar o raciocínio no documento correspondente.",
    ]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.18
        p.style = doc.styles["Normal"]

    # Section 2
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run("2. Estrutura padrão para novos documentos")
    format_run(r, size=17, bold=True, color="123E4A")

    structure_table = doc.add_table(rows=1, cols=2)
    structure_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    structure_table.autofit = False
    widths = [Cm(5.2), Cm(11.0)]
    headers = ["Elemento", "O que deve conter"]
    values = [
        ("Contexto", "Problema que motivou o documento, cenário atual e por que o tema importa para o Alfred."),
        ("Decisões", "Escolhas tomadas, alternativas descartadas e justificativas objetivas."),
        ("Impactos", "Efeitos em produto, dados, operação, experiência do usuário e riscos."),
        ("Regras e exceções", "Casos especiais, limites, dependências e comportamentos fora do fluxo padrão."),
        ("Referências", "Links, fontes, telas, regras do negócio, legislação, métricas ou trechos de código relacionados."),
        ("Próximos passos", "Itens pendentes, responsáveis e pontos que ainda precisam de validação."),
    ]

    hdr = structure_table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, "123E4A")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        format_run(r, size=10.5, bold=True, color="FFFFFF")

    for label, desc in values:
        row = structure_table.add_row()
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        for idx, content in enumerate([label, desc]):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(content)
            format_run(r, size=10.2, bold=(idx == 0), color="203040" if idx else "123E4A")
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row in structure_table.rows[1:]:
        for cell in row.cells:
            set_cell_shading(cell, "F8FBFC")

    # Section 3
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run("3. Catálogo inicial de documentos")
    format_run(r, size=17, bold=True, color="123E4A")

    catalog = doc.add_table(rows=1, cols=4)
    catalog.alignment = WD_TABLE_ALIGNMENT.CENTER
    catalog.autofit = False
    cat_widths = [Cm(2.8), Cm(4.2), Cm(4.9), Cm(4.3)]
    cat_headers = ["Tema", "Foco principal", "Perguntas centrais", "Entregáveis sugeridos"]
    items = [
        (
            "Design",
            "Interface, linguagem visual, hierarquia, legibilidade, fluxos e acessibilidade.",
            "Como o Alfred deve parecer e se comportar? Quais telas precisam de padronização?",
            "Guia visual, sistema de componentes, princípios de UX, checklist de acessibilidade.",
        ),
        (
            "LGPD / Regulamentação",
            "Tratamento de dados pessoais, base legal, retenção, consentimento e governança.",
            "Quais dados são coletados? Onde são armazenados? Por quanto tempo? Quem acessa?",
            "Mapa de dados, política de retenção, checklist de conformidade, análise de risco.",
        ),
        (
            "Engenharia de Dados",
            "Ingestão, modelagem, qualidade, integração com fontes e rastreabilidade.",
            "Como os dados entram no sistema? Como validamos, versionamos e recuperamos falhas?",
            "Arquitetura de dados, dicionário técnico, linhagem, runbook operacional.",
        ),
        (
            "Ciência de Dados",
            "Métricas, análises, previsões, segmentações e insights acionáveis.",
            "Quais perguntas o produto deve responder com dados? Como medir valor e comportamento?",
            "Catálogo de métricas, relatórios analíticos, notebooks, hipóteses e experimentos.",
        ),
        (
            "Contabilidade",
            "Classificação financeira, conciliações, fluxo de caixa e visão patrimonial.",
            "Como os movimentos devem ser categorizados? Como garantir consistência contábil?",
            "Plano de contas, regras de classificação, reconciliação, glossário financeiro.",
        ),
    ]

    hdr = catalog.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(cat_headers):
        cell = hdr.cells[i]
        cell.width = cat_widths[i]
        set_cell_shading(cell, "1F6F78")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        format_run(r, size=10, bold=True, color="FFFFFF")

    for row_data in items:
        row = catalog.add_row()
        for i, content in enumerate(row_data):
            cell = row.cells[i]
            cell.width = cat_widths[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(content)
            if i == 0:
                format_run(r, size=10.2, bold=True, color="123E4A")
            else:
                format_run(r, size=9.6, color="203040")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for idx, row in enumerate(catalog.rows[1:], start=1):
        fill = "FAFCFD" if idx % 2 else "F3F8FA"
        for cell in row.cells:
            set_cell_shading(cell, fill)

    # Section 4
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run("4. Ordem sugerida para produzir os próximos arquivos")
    format_run(r, size=17, bold=True, color="123E4A")

    order = [
        "1. Design - define linguagem, hierarquia visual e experiência base do produto.",
        "2. LGPD / Regulamentação - fixa limites e cuidados para o uso de dados.",
        "3. Engenharia de Dados - registra a arquitetura e a qualidade da base.",
        "4. Ciência de Dados - consolida métricas, análises e hipóteses de evolução.",
        "5. Contabilidade - fecha a visão financeira com regras de classificação e conciliação.",
    ]
    for item in order:
        numero, texto = item.split(" - ", 1)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(numero + " ")
        format_run(r, size=11, bold=True, color="123E4A")
        r = p.add_run(texto)
        format_run(r, size=11, color="203040")

    # Closing callout
    doc.add_paragraph()
    add_shaded_callout(
        doc,
        "Próximo passo recomendado",
        [
            "Criar um documento por tema a partir deste mapa, começando por Design e LGPD/Regulamentação.",
            "Depois, usar Engenharia de Dados e Ciência de Dados para consolidar a base técnica e analítica.",
            "Por fim, formalizar Contabilidade para fechar as regras financeiras e evitar interpretações divergentes.",
        ],
        fill="EEF6F3",
    )

    return doc


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_FILE)
    print(str(OUTPUT_FILE))


if __name__ == "__main__":
    main()
